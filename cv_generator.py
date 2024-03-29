import os
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
import PIL.Image

# paste the path where you installed the files
project_path = r"C:\Users\USer\OneDrive\Documents\Python\CV_Generator"
os.chdir(project_path)
 
document = Document()

def picture():
    path = input("Enter a valid pathname to an image:\n")
    try:
        image = PIL.Image.open(path)
        return image  # Return the image object
    except Exception as e:
        print(f"Error: {e}")
        return None  # Return None if there's an error

# Call the picture() function to get the image object
img = picture()

# If img is None (e.g., if there was an error), exit the program
if img is None:
    exit()

# Create a mask (circular)
mask = Image.new('L', img.size, 0)
draw = ImageDraw.Draw(mask)
draw.ellipse((0, 0, img.size[0], img.size[1]), fill=255)

# Apply the mask to the image
img = Image.composite(img, Image.new('RGBA', img.size, (255, 255, 255, 0)), mask)

# Save the cropped image
img.save('cropped_image.png')

# profile picture
picture = document.add_picture(
    'cropped_image.png', 
    width = Inches(2)
    )

# name phone number and email details
name = input('What is your full name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date: ')
to_date = input ('To Date: ')
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No: '
    )
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input ('To Date ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    elif has_more_experiences.lower() == 'no':
        break
    else:
        print('Please enter yes or no.')
        
# list of skills
document.add_heading('Skills')
skill = input('Enter a skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# more skills
while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No: '
    )
    if has_more_skills.lower() == 'yes':
        skill = input('Enter a skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break   
    
# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Edjenson's CV Generator using Python"

document.save('cv.docx')
